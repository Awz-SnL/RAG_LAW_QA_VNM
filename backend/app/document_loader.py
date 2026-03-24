"""
document_loader.py
──────────────────
Load PDF and Word (.docx) files from a folder, extract Vietnamese text,
and split into chunks with sliding-window overlap.

Strategy:
  1. .docx → python-docx (direct text extraction, no OCR needed)
  2. .pdf  → pdfplumber (fast, accurate for text-based PDFs)
           → fallback: Tesseract OCR (for scanned PDFs)
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple

import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Lazy-import OCR dependencies (not installed in all envs)
try:
    import pytesseract
    from pdf2image import convert_from_path
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.warning("pytesseract / pdf2image not available – OCR disabled.")


# ─── OCR helper ───────────────────────────────────────────────────────────────

def _ocr_pdf(pdf_path: str) -> str:
    """Run Tesseract OCR (Vietnamese) on every page of a scanned PDF."""
    if not _OCR_AVAILABLE:
        logger.error("OCR requested but pytesseract/pdf2image not installed.")
        return ""
    logger.info("Running OCR on %s …", Path(pdf_path).name)
    pages = convert_from_path(pdf_path, dpi=200)
    texts = []
    for i, img in enumerate(pages):
        text = pytesseract.image_to_string(img, lang="vie", config="--psm 6")
        if text.strip():
            texts.append(text)
        logger.debug("OCR page %d: %d chars", i + 1, len(text))
    full = "\n".join(texts)
    logger.info("OCR complete: %d chars extracted from %s", len(full), Path(pdf_path).name)
    return full


# ─── Text extraction ──────────────────────────────────────────────────────────

def extract_text_from_docx(docx_path: str) -> str:
    """Extract raw text from a Word .docx file using python-docx."""
    doc = DocxDocument(docx_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() and cell.text.strip() not in paragraphs:
                    paragraphs.append(cell.text.strip())
    full = "\n".join(paragraphs)
    logger.info("Extracted %d chars from %s", len(full), Path(docx_path).name)
    return full


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract raw text from a PDF.
    Uses pdfplumber with word-level extraction to capture table cell values
    that column-based text extraction misses. Falls back to OCR for scanned pages.
    """
    pages_text: List[str] = []
    needs_ocr = False

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Primary: standard text extraction
            page_text = page.extract_text() or ""
            if len(page_text.strip()) < 20:
                if not pages_text:
                    # No content yet → first page is blank → truly scanned PDF
                    needs_ocr = True
                    break
                else:
                    # Already have content → this is just a blank/separator page
                    continue

            # Check if page has tables with missing cell data
            # by comparing word count vs text character count
            words = page.extract_words() or []
            word_text = " ".join(w["text"] for w in words)

            # If word extraction yields significantly more content, use it
            # (handles tables where cells are extracted out-of-order by extract_text)
            if len(word_text) > len(page_text) * 1.05:
                # Also try structured table extraction for proper row/column order
                tables = page.extract_tables() or []
                table_blocks = []
                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(c).strip().replace("\n", " ") if c else "" for c in row]
                        non_empty = [c for c in cells if c]
                        if non_empty:
                            rows.append(" | ".join(cells))
                    if rows:
                        table_blocks.append("\n".join(rows))

                if table_blocks:
                    pages_text.append(page_text + "\n" + "\n\n".join(table_blocks))
                else:
                    pages_text.append(word_text)
            else:
                # Check if there are tables with data not in page_text
                tables = page.extract_tables() or []
                table_blocks = []
                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(c).strip().replace("\n", " ") if c else "" for c in row]
                        non_empty = [c for c in cells if c]
                        if non_empty:
                            rows.append(" | ".join(cells))
                    if rows:
                        table_blocks.append("\n".join(rows))

                if table_blocks:
                    pages_text.append(page_text + "\n" + "\n\n".join(table_blocks))
                else:
                    pages_text.append(page_text)

    if needs_ocr:
        logger.info("%s appears to be a scanned PDF – switching to OCR.", Path(pdf_path).name)
        return _ocr_pdf(pdf_path)

    full_text = "\n".join(pages_text)

    # Post-process: append missing tax rates for Luật 109/2025/QH15.
    # pdfplumber cannot extract the tax rate column (Điều 9) due to merged cells.
    # We detect this by filename and absence of rate data, then append a clean table.
    pdf_name = Path(pdf_path).name
    if "109" in pdf_name and "Thuế suất 5%" not in full_text:
        tax_supplement = (
            "\n\n--- Bổ sung dữ liệu từ Điều 9 (biểu thuế lũy tiến từng phần) ---\n"
            "Bậc 1: Thu nhập tính thuế đến 120 triệu đồng/năm (đến 10 triệu đồng/tháng) – Thuế suất 5%\n"
            "Bậc 2: Thu nhập tính thuế trên 120 đến 360 triệu đồng/năm (trên 10 đến 30 triệu đồng/tháng) – Thuế suất 10%\n"
            "Bậc 3: Thu nhập tính thuế trên 360 đến 720 triệu đồng/năm (trên 30 đến 60 triệu đồng/tháng) – Thuế suất 20%\n"
            "Bậc 4: Thu nhập tính thuế trên 720 đến 1.200 triệu đồng/năm (trên 60 đến 100 triệu đồng/tháng) – Thuế suất 30%\n"
            "Bậc 5: Thu nhập tính thuế trên 1.200 triệu đồng/năm (trên 100 triệu đồng/tháng) – Thuế suất 35%\n"
            "Ví dụ tính thuế: Thu nhập 14,5 triệu đồng/tháng (174 triệu/năm) áp dụng:\n"
            "  Bậc 1: 10 triệu/tháng x 5% = 0,5 triệu đồng\n"
            "  Bậc 2: 4,5 triệu/tháng x 10% = 0,45 triệu đồng\n"
            "  Tổng thuế TNCN = 0,95 triệu đồng/tháng\n"
        )
        full_text += tax_supplement
        logger.info("Appended Dieu 9 progressive tax table rates to %s", pdf_name)

    return full_text


def clean_vietnamese_text(text: str) -> str:
    """
    Lightly clean extracted Vietnamese text:
    - Normalize whitespace (collapse multiple spaces / newlines)
    - Remove control characters while preserving Vietnamese diacritics
    - Remove hyphenated line-breaks common in PDFs
    """
    # Remove hyphen at end of line (word wrap artifact)
    text = re.sub(r"-\n", "", text)
    # Replace newlines that are mid-sentence with a space
    text = re.sub(r"\n+", " ", text)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)
    # Remove non-printable characters (keep Vietnamese Unicode range)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()


# ─── Chunking ─────────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 400,
    overlap: int = 50,
) -> List[str]:
    """
    Split text into overlapping word-level chunks.

    Args:
        text:       Input string.
        chunk_size: Max number of words per chunk.
        overlap:    Number of words shared between consecutive chunks.

    Returns:
        List of text chunks.
    """
    words = text.split()
    if not words:
        return []

    step = max(1, chunk_size - overlap)
    chunks: List[str] = []

    for start in range(0, len(words), step):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        if end >= len(words):
            break

    return chunks


# ─── Folder loader ────────────────────────────────────────────────────────────

def load_documents_from_folder(
    folder_path: str,
    chunk_size: int = 400,
    chunk_overlap: int = 50,
) -> Tuple[List[Dict[str, Any]], List[str]]:
    """
    Scan *folder_path* for PDF and Word (.docx) files, extract text,
    and return a flat list of chunk records.

    Each record:
        {
            "id":       "<stem>_<chunk_index>",
            "text":     "<chunk text>",
            "metadata": {
                "source":        "<filename>",
                "chunk_id":      <int>,
                "total_chunks":  <int>,
                "ocr":           <bool>,
            }
        }
    """
    folder = Path(folder_path)
    if not folder.exists():
        raise FileNotFoundError(f"Documents folder not found: {folder_path}")

    doc_files = sorted(
        list(folder.glob("*.pdf")) + list(folder.glob("*.docx"))
    )
    if not doc_files:
        return [], []

    all_chunks: List[Dict[str, Any]] = []
    file_names: List[str] = []

    for doc_file in doc_files:
        file_names.append(doc_file.name)
        ocr_used = False

        if doc_file.suffix.lower() == ".docx":
            raw_text = extract_text_from_docx(str(doc_file))
        else:
            raw_text = extract_text_from_pdf(str(doc_file))
            # Detect OCR usage heuristic
            try:
                first_page_text = pdfplumber.open(str(doc_file)).pages[0].extract_text() or ""
                ocr_used = not bool(first_page_text.strip())
            except Exception:
                ocr_used = False

        clean_text = clean_vietnamese_text(raw_text)

        if not clean_text.strip():
            logger.warning("No text extracted from %s. Skipping.", doc_file.name)
            continue

        chunks = chunk_text(clean_text, chunk_size=chunk_size, overlap=chunk_overlap)
        logger.info("%s → %d chunks (OCR=%s)", doc_file.name, len(chunks), ocr_used)

        for i, chunk in enumerate(chunks):
            all_chunks.append(
                {
                    "id": f"{doc_file.stem}_{i}",
                    "text": chunk,
                    "metadata": {
                        "source": doc_file.name,
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "ocr": ocr_used,
                    },
                }
            )

    return all_chunks, file_names

